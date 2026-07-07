#!/usr/bin/env python3
"""Build a SecureScan Word document from Markdown."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


ROOT_DIR = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT_DIR / "docs" / "SECURESCAN-OVERVIEW.md"
DEFAULT_OUTPUT = ROOT_DIR / "docs" / "SecureScan-Overview.docx"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a SecureScan Word document from Markdown.")
    parser.add_argument("--source", default=str(DEFAULT_SOURCE), help="Markdown source path.")
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT), help="DOCX output path.")
    parser.add_argument("--footer-title", default="SecureScan Overview", help="Footer title.")
    return parser.parse_args()


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for name, size in [("Title", 22), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True


def add_code_block(document: Document, lines: list[str]) -> None:
    if not lines:
        return
    paragraph = document.add_paragraph()
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index in range(width):
            text = row[column_index] if column_index < len(row) else ""
            cell = table.cell(row_index, column_index)
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                    if row_index == 0:
                        run.font.bold = True


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        line = lines[index].strip()
        cells = [cell.strip().replace("`", "") for cell in line.strip("|").split("|")]
        is_separator = all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)
        if not is_separator:
            rows.append(cells)
        index += 1
    return rows, index


def add_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue

        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=0)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif re.match(r"\d+\. ", stripped):
            document.add_paragraph(re.sub(r"^\d+\. ", "", stripped).strip(), style="List Number")
        else:
            document.add_paragraph(stripped)
        index += 1

    add_code_block(document, code_lines)


def main() -> int:
    args = parse_args()
    source = Path(args.source).expanduser().resolve()
    output = Path(args.output).expanduser().resolve()

    document = Document()
    configure_document(document)
    add_markdown(document, source.read_text(encoding="utf-8"))
    footer = document.sections[0].footer.paragraphs[0]
    footer.text = args.footer_title
    footer.runs[0].font.size = Pt(8)
    footer.add_run().add_break(WD_BREAK.LINE)
    footer.add_run(f"Generated from {source.relative_to(ROOT_DIR).as_posix()}").font.size = Pt(8)
    document.save(output)
    print(f"Built {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
